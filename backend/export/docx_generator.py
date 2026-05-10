from docx import Document
from docx.shared import Pt, RGBColor
from io import BytesIO

def generate_docx(query: str, report: str) -> bytes:
    doc = Document()

    title = doc.add_heading("Research Agent Report", 0)
    title.runs[0].font.color.rgb = RGBColor(0x1D, 0x9E, 0x75)

    doc.add_paragraph(f"Query: {query}")
    doc.add_paragraph("")

    for line in report.split("\n"):
        stripped = line.strip()
        if stripped.startswith("## "):
            doc.add_heading(stripped[3:], level=2)
        elif stripped.startswith("- "):
            doc.add_paragraph(stripped[2:], style="List Bullet")
        elif stripped:
            doc.add_paragraph(stripped)

    buffer = BytesIO()
    doc.save(buffer)
    return buffer.getvalue()
